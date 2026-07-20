import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "Pydurma", "src"))
"""
Tibetan Collation Web App — Streamlit
======================================
Run with:
    pip install streamlit bayoo-docx
    streamlit run collation_app.py
"""

import io
import re
import tempfile
from pathlib import Path

import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from Pydurma.gen.normalizer_gen import GenericNormalizer
from Pydurma.gen.tokenizer_gen import GenericTokenizer
from Pydurma.encoder import Encoder
from Pydurma.aligners.fdmp import FDMPaligner
from Pydurma.utils.utils import column_matrix_to_row_matrix, token_row_to_text_row


# ─────────────────────────────────────────────
#  CONSTANTS
# ─────────────────────────────────────────────

COLOR_LIST = [
    "FFD37F", "9AD1FF", "A5FFB1", "FF9A9A", "C79CFF",
    "FFB27F", "7AFFD5", "FF7FBA", "B5FF7A", "7FB1FF",
    "FF7F7F", "FFE07F",
]

# Shad (Tibetan sentence/clause punctuation) — kept separate so it can be
# toggled on/off during collation. U+0F0D..U+0F11 and U+0F14. The ASCII "/"
# is the Wylie/EWTS transliteration of the shad, so transliterated texts get
# the same shad-ignore treatment (a single "/" and a double "//" both reduce
# to shad characters that are stripped when shad differences are ignored).
SHAD_CHARS_CORE = set(["།", "༎", "༏", "༐", "༑", "༔", "/"])

# Characters that are always ignored when detecting content differences
# (tsheg, yig-mgo, brackets, spaces, western punctuation).
PUNCT_TO_IGNORE_CORE = set([
    "་", "༄", "༅", "༈",
    "༼", "༽", "༌", "༗", "༘",
    " ", "\n", "\t", ",", ".", "?", ":", ";",
])

# Effective sets — (re)built by apply_preprocessing_options() below. The UI
# calls it with the user's preprocessing choices before running a collation;
# the module-level call further down installs the defaults for library use.
SHAD_CHARS = set(SHAD_CHARS_CORE)
PUNCT_TO_IGNORE_BASE = set(PUNCT_TO_IGNORE_CORE)
PUNCT_TO_IGNORE = PUNCT_TO_IGNORE_BASE | SHAD_CHARS


# ── Preprocessing (input normalization) ──────────────────────────────
# Folio/page tags like [354], [zhe 1], [kha zhe lnga] are reference markers,
# not text; left in place they collate as readings and pollute the apparatus.
_FOLIO_TAG_RE = re.compile(r"\[[^\[\]\n]{1,40}\]")


def strip_folio_tags(text: str) -> str:
    """Remove [..] folio/page tags, leaving a space so words don't fuse."""
    return _FOLIO_TAG_RE.sub(" ", text)


def extract_folio_tags(text: str):
    """Strip [..] tags like strip_folio_tags, but remember what and where.

    Returns ``(stripped_text, milestones)`` where milestones is a list of
    ``(offset, tag)`` pairs: ``offset`` is the character position in the
    *stripped* text (pointing at the space that replaced the tag) where the
    tag can be re-inserted later, e.g. into the golden document.
    """
    milestones = []
    out = []
    pos = 0  # length of stripped text built so far
    last = 0
    for m in _FOLIO_TAG_RE.finditer(text):
        out.append(text[last : m.start()])
        pos += m.start() - last
        milestones.append((pos, m.group(0)))
        out.append(" ")
        pos += 1
        last = m.end()
    out.append(text[last:])
    return "".join(out), milestones


def count_preprocessing_hits(text: str) -> dict:
    """Per-rule occurrence counts, for the preprocessing preview."""
    return {
        "folio/page tags [..]": len(_FOLIO_TAG_RE.findall(text)),
        "underscores _": text.count("_"),
        "pipes |": text.count("|"),
        "head marks @ # !": sum(text.count(c) for c in "@#!"),
    }


def apply_preprocessing_options(
    underscore_as_space=True,
    pipe_as_shad=True,
    ignore_head_marks=True,
):
    """Install the effective character sets used by the collation.

    - underscore_as_space: EWTS writes an explicit space as ``_``; treat it
      as whitespace everywhere (ignore set, syllable splitting, a-chung
      reattachment).
    - pipe_as_shad: ``|`` is an alternate EWTS shad (common in OCR output).
    - ignore_head_marks: ``@``, ``#``, ``!`` transliterate yig-mgo ornaments
      (༄༅ …), which are structural, not textual.

    Rebinding the module-level sets keeps every existing function signature
    unchanged; the display/golden text itself is never rewritten by these.
    """
    global SHAD_CHARS, PUNCT_TO_IGNORE_BASE, PUNCT_TO_IGNORE
    global _SYLLABLE_SEP_RE, _SEP_CHARS
    SHAD_CHARS = set(SHAD_CHARS_CORE) | ({"|"} if pipe_as_shad else set())
    PUNCT_TO_IGNORE_BASE = (
        set(PUNCT_TO_IGNORE_CORE)
        | ({"_"} if underscore_as_space else set())
        | ({"@", "#", "!"} if ignore_head_marks else set())
    )
    PUNCT_TO_IGNORE = PUNCT_TO_IGNORE_BASE | SHAD_CHARS
    _SYLLABLE_SEP_RE = re.compile(r"[་༌\s_]+" if underscore_as_space else r"[་༌\s]+")
    _SEP_CHARS = set(["་", "༌", " ", "\t", "\n"]) | (
        {"_"} if underscore_as_space else set()
    )


# ─────────────────────────────────────────────
#  CORE LOGIC (unchanged from your script)
# ─────────────────────────────────────────────

def ensure_footnote_reference_style(document):
    """Define the FootnoteReference character style with superscript.

    bayoo-docx emits each in-text reference mark as
    <w:rStyle w:val="FootnoteReference"/> but never defines that style, so
    Word renders the number at the baseline. We create it here so the mark
    renders raised/superscript like a real footnote number.
    """
    styles_el = document.styles.element
    target = None
    for st_el in styles_el.findall(qn("w:style")):
        if st_el.get(qn("w:styleId")) == "FootnoteReference":
            target = st_el
            break
    if target is None:
        target = OxmlElement("w:style")
        target.set(qn("w:type"), "character")
        target.set(qn("w:styleId"), "FootnoteReference")
        name = OxmlElement("w:name")
        name.set(qn("w:val"), "footnote reference")
        target.append(name)
        styles_el.append(target)
    rPr = target.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        target.append(rPr)
    va = rPr.find(qn("w:vertAlign"))
    if va is None:
        va = OxmlElement("w:vertAlign")
        rPr.append(va)
    va.set(qn("w:val"), "superscript")


def shrink_footnote_style(document, font_size=8, line_spacing_multiple=0.85):
    styles = document.styles
    for style_name in ("Footnote Text", "Footnote Reference"):
        try:
            s = styles[style_name]
            s.font.size = Pt(font_size)
            if style_name == "Footnote Reference":
                # ensure the in-text reference mark renders raised/superscript
                s.font.superscript = True
            if style_name == "Footnote Text":
                pf = s.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = line_spacing_multiple
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
        except KeyError:
            pass


def strip_ignorable(s: str, ignore_shad: bool = True) -> str:
    """Remove characters that don't count as content differences.

    When ignore_shad is True (default), shad punctuation is also stripped, so
    shad-only differences won't generate notes. When False, shad is preserved
    and therefore shad differences will surface as variant notes.
    """
    ignore_set = PUNCT_TO_IGNORE if ignore_shad else PUNCT_TO_IGNORE_BASE
    if not ignore_shad and "|" in SHAD_CHARS:
        # "|" and "/" are the same shad in different notation; when shad is
        # kept for comparison they must not read as a difference.
        s = s.replace("|", "/")
    return "".join(ch for ch in s if ch not in ignore_set)


_SEP_CHARS = set(["་", "༌", " ", "\t", "\n"])


def _reattach_stranded_achung(*rows):
    """Move a lone initial a-chung onto the syllable that follows it.

    The aligner sometimes strands an initial a-chung (Wylie ``'`` / Unicode
    ``འ``) at the end of one cell while the syllable it belongs to lands in the
    next cell — e.g. witness ``"su '"`` + ``"gyur"`` against base ``"su "`` +
    ``"gyur"``. Left alone this reads as a spurious ``su'`` variant. Moving the
    a-chung forward yields base ``"gyur"`` vs witness ``"'gyur"`` so the note
    correctly reads ``gyur] 'gyur``. Only a *lone* a-chung (preceded by a
    separator, i.e. an initial one) is moved; a final a-chung glued to letters
    such as ``dga'`` is left untouched. Per-row concatenation is preserved.
    """
    for row in rows:
        if row is None:
            continue
        n = len(row)
        for i in range(n):
            seg = row[i]
            if not seg or seg == "-":
                continue
            # Only when the a-chung is the very last character of the cell can
            # it move to the next cell's front without reordering anything in
            # between, so the base/golden text stays byte-for-byte intact.
            if seg[-1] not in A_CHUNG_CHARS:
                continue
            k = len(seg)
            while k > 0 and seg[k - 1] in A_CHUNG_CHARS:
                k -= 1
            # Lone/initial a-chung only: preceded by a separator, a shad, or
            # nothing. After a shad a new word begins, so a lone a-chung there
            # can only be the initial letter of the next word (e.g. "/ /'"
            # before "dir" is the 'a of "'dir").
            if k > 0 and seg[k - 1] not in _SEP_CHARS and seg[k - 1] not in SHAD_CHARS:
                continue
            run = seg[k:]
            j = i + 1
            while j < n and (not row[j] or row[j] == "-"):
                j += 1
            if j >= n:
                continue
            row[i] = seg[:k]
            row[j] = run + row[j]
    return rows


def align_three(text1: str, text2: str, text3: str):
    normalizer = GenericNormalizer()
    encoder = Encoder()
    tokenizer = GenericTokenizer(encoder, normalizer)
    aligner = FDMPaligner()

    tokens1, tokenstr1 = tokenizer.tokenize(text1)
    tokens2, tokenstr2 = tokenizer.tokenize(text2)
    tokens3, tokenstr3 = tokenizer.tokenize(text3)

    matrix = aligner.get_alignment_matrix(
        [tokenstr1, tokenstr2, tokenstr3],
        [tokens1, tokens2, tokens3],
    )
    row_matrix = column_matrix_to_row_matrix(matrix)

    aligned1 = token_row_to_text_row(row_matrix[0], text1)
    aligned2 = token_row_to_text_row(row_matrix[1], text2)
    aligned3 = token_row_to_text_row(row_matrix[2], text3)
    _reattach_stranded_achung(aligned1, aligned2, aligned3)
    return aligned1, aligned2, aligned3


def align_two(text1: str, text2: str):
    """2-way alignment. Returns aligned3 as None to signal single-comparison mode."""
    normalizer = GenericNormalizer()
    encoder = Encoder()
    tokenizer = GenericTokenizer(encoder, normalizer)
    aligner = FDMPaligner()

    tokens1, tokenstr1 = tokenizer.tokenize(text1)
    tokens2, tokenstr2 = tokenizer.tokenize(text2)

    matrix = aligner.get_alignment_matrix(
        [tokenstr1, tokenstr2],
        [tokens1, tokens2],
    )
    row_matrix = column_matrix_to_row_matrix(matrix)

    aligned1 = token_row_to_text_row(row_matrix[0], text1)
    aligned2 = token_row_to_text_row(row_matrix[1], text2)
    _reattach_stranded_achung(aligned1, aligned2)
    # None signals to export functions that there is no third version
    return aligned1, aligned2, None


def set_run_background_color(run, hex_color: str):
    hex_color = hex_color.lstrip("#").upper()
    rPr = run._element.get_or_add_rPr()
    shd = rPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        rPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


OMITTED_MARK = "om."  # standard critical-apparatus mark for an omitted reading

TSHEG = "་"  # Tibetan intersyllabic tsheg, used to rejoin syllables

# a-chung: U+0F60 (འ) in Unicode Tibetan, apostrophe (') in Wylie/EWTS. When an
# alignment boundary strands a bare a-chung as its own token, it should re-join
# the adjacent syllable so an added/omitted a-chung reads as e.g. "su'" rather
# than surfacing as a meaningless standalone "'".
A_CHUNG_CHARS = set(["འ", "'"])

# Syllable separators for display splitting: tsheg, no-break tsheg, whitespace.
_SYLLABLE_SEP_RE = re.compile(r"[་༌\s]+")


def _syllables(seg: str, ignore_shad: bool):
    """Split an aligned segment into display syllables.

    Tsheg and whitespace act as separators (and are dropped). Shad is removed
    only when shad differences are ignored, so it stays visible otherwise.
    Letters (and any residual marks) are preserved inside each syllable. A bare
    a-chung is re-attached to its neighbour rather than kept as its own token.
    """
    if not seg:
        return []
    s = seg
    if ignore_shad:
        for ch in SHAD_CHARS:
            s = s.replace(ch, "")
    # Drop ignorable non-content characters (head marks, western punctuation)
    # from the display so readings never show e.g. "@#" from a source file.
    parts = []
    for p in _SYLLABLE_SEP_RE.split(s):
        p = "".join(c for c in p if c not in PUNCT_TO_IGNORE_BASE)
        if p:
            parts.append(p)
    merged = []
    pending = ""  # a leading bare a-chung waiting to attach to the next syllable
    for p in parts:
        if all(c in A_CHUNG_CHARS for c in p):
            if merged:
                merged[-1] = merged[-1] + p
            else:
                pending += p
        else:
            merged.append(pending + p)
            pending = ""
    if pending:
        merged.append(pending)
    return merged


def _trim_common_syllables(readings):
    """Trim syllables shared by *all* readings at the start and the end.

    ``readings`` is a list of syllable-lists. Returns new syllable-lists with
    the common leading/trailing syllables removed, so only the differing part
    remains. Because at least one reading always differs when a note exists,
    this never trims a reading down to nothing on every side simultaneously.
    """
    readings = [list(r) for r in readings]
    if len(readings) < 2:
        return readings
    # common prefix
    n = min(len(r) for r in readings)
    pre = 0
    while pre < n and all(r[pre] == readings[0][pre] for r in readings):
        pre += 1
    readings = [r[pre:] for r in readings]
    # common suffix
    n = min(len(r) for r in readings)
    suf = 0
    while suf < n and all(r[-1 - suf] == readings[0][-1 - suf] for r in readings):
        suf += 1
    if suf:
        readings = [r[: len(r) - suf] for r in readings]
    return readings


_TIBETAN_CHAR_RE = re.compile(r"[ༀ-࿿]")


def _reading_display(sylls) -> str:
    """Render a (trimmed) syllable-list for a note; empty = an omission.

    Syllables are rejoined with a tsheg for Tibetan-script input but with a
    plain space for Wylie/roman input — a tsheg between roman letters would
    mix scripts (e.g. "sgrog་la'ang" instead of "sgrog la'ang").
    """
    if not sylls:
        return OMITTED_MARK
    joiner = TSHEG if any(_TIBETAN_CHAR_RE.search(s) for s in sylls) else " "
    return joiner.join(sylls)


def build_note_text(
    seg1, seg2, seg3,
    two_way=False,
    positive=False,
    ignore_shad=True,
    label1="V1", label2="V2", label3="V3",
) -> str:
    """Build a single apparatus note in classic critical-edition style.

    Format: ``<baseSiglum> <lemma>] <sigla> <reading>; <sigla> <reading>``

    - The lemma is the base/golden reading (or ``om.`` when the base omits it),
      and is itself labelled with the base siglum so it can be moved into the
      variant list unchanged if the base is later reassigned.
    - Sigla precede the reading they belong to on both sides of the bracket.
      Sigla that share a reading are comma-separated (``AB1, GB1``); distinct
      readings are separated by ``; ``. No colon is used.
    - Multi-syllable segments are reduced to just the differing syllable(s):
      syllables shared by every witness are trimmed away, and the surviving
      syllables keep their tsheg separators so they read correctly.
    - Negative apparatus (default): only witnesses that differ from the lemma
      are listed. A positive apparatus additionally credits the witnesses
      that agree with the base by listing their sigla with the lemma
      (``BX1, AB1 la] GB1 pa``) instead of repeating the reading.

    Witnesses sharing the same reading are grouped, e.g. ``GX1, GB1 ...``.
    """
    # comparison keys (punctuation/tsheg-insensitive) decide agreement
    key1 = strip_ignorable(seg1, ignore_shad)
    key2 = strip_ignorable(seg2, ignore_shad)
    key3 = "" if two_way else strip_ignorable(seg3, ignore_shad)

    # display syllable lists (tsheg preserved between syllables)
    s1 = _syllables(seg1, ignore_shad)
    s2 = _syllables(seg2, ignore_shad)
    s3 = [] if two_way else _syllables(seg3, ignore_shad)

    # isolate the differing syllable(s) by trimming shared context
    all_readings = [s1, s2] if two_way else [s1, s2, s3]
    trimmed = _trim_common_syllables(all_readings)
    if two_way:
        t1, t2 = trimmed
        t3 = []
    else:
        t1, t2, t3 = trimmed

    lemma = _reading_display(t1)

    # (label, comparison_key, trimmed_syllables) per comparison witness
    witnesses = [(label2, key2, t2)]
    if not two_way:
        witnesses.append((label3, key3, t3))

    # Witnesses agreeing with the base support the lemma. A positive
    # apparatus records that by listing their sigla with the lemma —
    # "BX1, AB1 la] GB1 pa" — rather than repeating the reading as a
    # variant. A negative apparatus leaves them out entirely.
    lemma_labels = [label1]
    if positive:
        lemma_labels += [lab for (lab, k, t) in witnesses if k == key1]

    selected = [(lab, t) for (lab, k, t) in witnesses if k != key1]

    if not selected:
        return ""

    # group witnesses that share the same displayed reading, preserving order
    groups = []  # list of [reading_display, [labels...]]
    for lab, t in selected:
        disp = _reading_display(t)
        for g in groups:
            if g[0] == disp:
                g[1].append(lab)
                break
        else:
            groups.append([disp, [lab]])

    parts = [f"{', '.join(labs)} {disp}" for disp, labs in groups]
    return f"{', '.join(lemma_labels)} {lemma}] " + "; ".join(parts)


def export_three_way_with_notes(
    aligned1, aligned2, aligned3,
    label1, label2, label3,
    name1="base", name2="comp1", name3="comp2",
    ignore_shad=True,
    positive=False,
):
    two_way = aligned3 is None  # single-comparison mode

    doc = Document()
    doc.add_heading("Tibetan Collation Report", level=1)
    if two_way:
        doc.add_paragraph(f"Base / golden: {name1}  |  Comparison: {name2}")
    else:
        doc.add_paragraph(f"Base / golden: {name1}  |  Comparison 1: {name2}  |  Comparison 2: {name3}")

    num_cols = 2 if two_way else 3
    table = doc.add_table(rows=2, cols=num_cols)
    hdr = table.rows[0].cells
    hdr[0].text = f"{label1} (golden, notes)"
    hdr[1].text = label2
    if not two_way:
        hdr[2].text = label3

    row = table.rows[1].cells
    p_v1 = row[0].paragraphs[0]
    p_v2 = row[1].paragraphs[0]
    p_v3 = row[2].paragraphs[0] if not two_way else None

    if two_way:
        aligned3 = []  # empty — never iterated directly
    max_len = max(len(aligned1), len(aligned2), len(aligned3) if aligned3 else 0)
    a1 = list(aligned1) + [""] * (max_len - len(aligned1))
    a2 = list(aligned2) + [""] * (max_len - len(aligned2))
    a3 = list(aligned3) + [""] * (max_len - len(aligned3)) if not two_way else [""] * max_len

    notes = []
    note_active = False
    current_note_color = None
    color_idx = -1

    for seg1, seg2, seg3 in zip(a1, a2, a3):
        seg2_raw = seg2
        seg3_raw = seg3 if not two_way else ""
        seg1 = "" if seg1 == "-" else seg1
        seg2 = "" if seg2 == "-" else seg2
        seg3 = "" if (two_way or seg3 == "-") else seg3

        norm1 = strip_ignorable(seg1, ignore_shad)
        norm2 = strip_ignorable(seg2, ignore_shad)
        norm3 = "" if two_way else strip_ignorable(seg3, ignore_shad)

        v2_missing = (seg2_raw == "-" or (seg2 and norm2 == "")) and norm1 != ""
        v3_missing = False if two_way else ((seg3_raw == "-" or (seg3 and norm3 == "")) and norm1 != "")
        v1_missing = (norm1 == "" and (norm2 != "" or (not two_way and norm3 != "")))

        diff12 = True if v2_missing else (norm1 != norm2) if norm1 or norm2 else False
        diff13 = False if two_way else (True if v3_missing else (norm1 != norm3) if norm1 or norm3 else False)

        has_real_diff = (norm1 != "" and (diff12 or diff13)) or v1_missing
        note_start_here = False

        if has_real_diff and not note_active:
            note_text = build_note_text(
                seg1, seg2, seg3,
                two_way=two_way, positive=positive, ignore_shad=ignore_shad,
                label1=label1, label2=label2, label3=label3,
            )
            if note_text:
                notes.append(note_text)
                note_start_here = True
                note_active = True
                color_idx = (color_idx + 1) % len(COLOR_LIST)
                current_note_color = COLOR_LIST[color_idx]

        color = current_note_color if note_active else None
        note_number = len(notes)

        if seg1:
            run1 = p_v1.add_run(seg1)
            if color and (diff12 or diff13) and norm1 != "":
                set_run_background_color(run1, color)
            if note_start_here:
                m = p_v1.add_run(f"[{note_number}]")
                m.font.superscript = True
        elif note_start_here:
            m = p_v1.add_run(f"[{note_number}]")
            m.font.superscript = True

        if seg2:
            run2 = p_v2.add_run(seg2)
            if color and (diff12 or v1_missing) and norm2 != "":
                set_run_background_color(run2, color)
            if note_start_here and (diff12 or v1_missing):
                m = p_v2.add_run(f"[{note_number}]")
                m.font.superscript = True

        if not two_way and p_v3 is not None and seg3:
            run3 = p_v3.add_run(seg3)
            if color and (diff13 or v1_missing) and norm3 != "":
                set_run_background_color(run3, color)
            if note_start_here and (diff13 or v1_missing):
                m = p_v3.add_run(f"[{note_number}]")
                m.font.superscript = True

        if note_active and not has_real_diff:
            note_active = False
            current_note_color = None

    if notes:
        doc.add_paragraph()
        doc.add_heading("Notes", level=2)
        for i, text in enumerate(notes, start=1):
            p = doc.add_paragraph()
            r = p.add_run(f"{i}")
            r.font.superscript = True
            r.bold = True
            p.add_run(" ")
            p.add_run(text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, notes


def _note_lemma_is_shad(note_text: str, label1: str) -> bool:
    """True when a note's lemma consists only of shad punctuation.

    Notes read ``<siglum> <lemma>] <variants>``; the lemma is what the
    footnote mark should sit on. When it is a shad, the mark must stay on
    the shad instead of moving back onto the preceding word.
    """
    close = note_text.find("]")
    if close <= 0:
        return False
    lemma = note_text[:close]
    prefix = label1 + " "
    if lemma.startswith(prefix):
        lemma = lemma[len(prefix):]
    lemma = lemma.strip()
    return bool(lemma) and all(c in SHAD_CHARS or c.isspace() for c in lemma)


def export_golden_with_footnotes(
    aligned1, aligned2, aligned3,
    notes,
    label1, label2, label3,
    name1="base",
    ignore_shad=True,
    milestones=None,
):
    """Golden text with variant footnotes.

    ``milestones`` is an optional list of ``(offset, tag)`` pairs from
    extract_folio_tags(): folio/page tags stripped before collation that are
    re-inserted here — as italic runs at their original character positions —
    without ever having been part of the alignment or the apparatus.
    """
    two_way = aligned3 is None

    milestones = sorted(milestones) if milestones else []
    ms_index = 0  # next milestone still to be emitted
    char_pos = 0  # running offset into the (concatenated) base text

    def emit(p, text):
        """Write base text, splicing in any milestone tags it spans."""
        nonlocal ms_index, char_pos
        start = 0
        end_pos = char_pos + len(text)
        while ms_index < len(milestones) and milestones[ms_index][0] < end_pos:
            cut = milestones[ms_index][0] - char_pos
            if cut > start:
                p.add_run(text[start:cut])
            tag_run = p.add_run(milestones[ms_index][1])
            tag_run.italic = True
            start = cut
            ms_index += 1
        if start < len(text):
            p.add_run(text[start:])
        char_pos = end_pos

    doc = Document()
    doc.add_heading(f"{label1} with Footnotes", level=1)
    if two_way:
        doc.add_paragraph(f"Base: {name1}  |  Footnotes from comparison with {label2}.")
    else:
        doc.add_paragraph(f"Base: {name1}  |  Footnotes from comparison with {label2} and {label3}.")
    doc.add_paragraph()

    p_text = doc.add_paragraph()

    if two_way:
        aligned3 = []
    max_len = max(len(aligned1), len(aligned2), len(aligned3) if aligned3 else 0)
    a1 = list(aligned1) + [""] * (max_len - len(aligned1))
    a2 = list(aligned2) + [""] * (max_len - len(aligned2))
    a3 = list(aligned3) + [""] * (max_len - len(aligned3)) if not two_way else [""] * max_len

    note_active = False
    note_index = 0

    for seg1, seg2, seg3 in zip(a1, a2, a3):
        seg2_raw = seg2
        seg3_raw = seg3 if not two_way else ""
        seg1 = "" if seg1 == "-" else seg1
        seg2 = "" if seg2 == "-" else seg2
        seg3 = "" if (two_way or seg3 == "-") else seg3

        norm1 = strip_ignorable(seg1, ignore_shad)
        norm2 = strip_ignorable(seg2, ignore_shad)
        norm3 = "" if two_way else strip_ignorable(seg3, ignore_shad)

        v2_missing = (seg2_raw == "-" or (seg2 and norm2 == "")) and norm1 != ""
        v3_missing = False if two_way else ((seg3_raw == "-" or (seg3 and norm3 == "")) and norm1 != "")
        v1_missing = (norm1 == "" and (norm2 != "" or (not two_way and norm3 != "")))

        diff12 = True if v2_missing else (norm1 != norm2) if norm1 or norm2 else False
        diff13 = False if two_way else (True if v3_missing else (norm1 != norm3) if norm1 or norm3 else False)

        has_real_diff = (norm1 != "" and (diff12 or diff13)) or v1_missing
        note_start_here = False

        if has_real_diff and not note_active:
            note_index += 1
            note_start_here = True
            note_active = True

        place_note = note_start_here and 1 <= note_index <= len(notes)

        if place_note and seg1:
            # Put the reference mark right after the annotated word, before any
            # trailing space or shad, so it renders as "su² gyur" not
            # "su ²gyur" and "grag go²/" not "grag go/²".
            #
            # Exception: when the lemma *is* a shad (only possible with shad
            # differences reported), the note is about that shad, so the mark
            # must stay on it rather than jump back to the preceding word.
            # The test has to use the note's lemma rather than the raw
            # segment: a segment often holds a word and a shad together
            # ("cing / "), while the lemma is trimmed down to just "/".
            lemma_is_shad = _note_lemma_is_shad(notes[note_index - 1], label1)
            cut = len(seg1)
            if not lemma_is_shad:
                while cut > 0 and (seg1[cut - 1].isspace() or seg1[cut - 1] in SHAD_CHARS
                                   or seg1[cut - 1] in PUNCT_TO_IGNORE_BASE):
                    cut -= 1
            else:
                while cut > 0 and seg1[cut - 1].isspace():
                    cut -= 1
            if cut == 0:  # nothing to anchor to — keep the original placement
                cut = len(seg1.rstrip())
            content = seg1[:cut]
            trailing = seg1[cut:]
            if content:
                emit(p_text, content)
            p_text.add_footnote(notes[note_index - 1])
            if trailing:
                emit(p_text, trailing)
        else:
            if seg1:
                emit(p_text, seg1)
            if place_note:
                p_text.add_footnote(notes[note_index - 1])

        if note_active and not has_real_diff:
            note_active = False

    # Any milestone past the last emitted character (e.g. a tag at the very
    # end of the base text) still needs to be written out.
    while ms_index < len(milestones):
        tag_run = p_text.add_run(milestones[ms_index][1])
        tag_run.italic = True
        ms_index += 1

    # Footnote styles are only present after add_footnote() has run, so apply
    # formatting now (before save) rather than on the empty document.
    ensure_footnote_reference_style(doc)
    shrink_footnote_style(doc, font_size=8, line_spacing_multiple=0.85)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# Install the default preprocessing sets (must run after every dependent
# constant above is defined; the UI re-applies with the user's choices).
apply_preprocessing_options()


# ─────────────────────────────────────────────
#  STREAMLIT UI
# ─────────────────────────────────────────────

st.set_page_config(
    page_title="Tibetan Collation Tool",
    page_icon="📜",
    layout="centered",
)

st.title("📜 Tibetan Collation Tool")
st.caption("Upload your texts, run the collation, and download the Word outputs.")

st.divider()

# ── File uploads
st.subheader("1 · Upload texts")

col_a, col_b = st.columns([1, 1])
with col_a:
    base_file = st.file_uploader(
        "Base / golden text (.txt)",
        type=["txt"],
        help="This is the primary version — all notes are anchored here.",
    )
    label1 = st.text_input("Label for base text", value="BX1")

with col_b:
    comp_mode = st.radio(
        "Number of comparison texts",
        options=["1 comparison text", "2 comparison texts"],
        index=1,
    )

st.divider()
st.subheader("2 · Comparison text(s)")

two_texts = comp_mode == "2 comparison texts"

col1, col2 = st.columns(2) if two_texts else (st.columns(1)[0], None)

with col1:
    comp1_file = st.file_uploader("Comparison text 1 (.txt)", type=["txt"], key="c1")
    label2 = st.text_input("Label for comparison 1", value="GX1")

if two_texts and col2 is not None:
    with col2:
        comp2_file = st.file_uploader("Comparison text 2 (.txt)", type=["txt"], key="c2")
        label3 = st.text_input("Label for comparison 2", value="GB1")
else:
    comp2_file = None
    label3 = "—"

st.divider()
st.subheader("3 · Options")

with st.expander("Preprocessing", expanded=True):
    st.caption(
        "Cleanup applied before collation. Each option only takes effect if "
        "its pattern actually appears in your files — otherwise it changes "
        "nothing. Leaving them all checked is safe."
    )
    prep_tags = st.checkbox(
        "Strip folio/page tags like [354], [zhe 1]",
        value=True,
        help="Bracketed reference markers are removed before alignment so "
        "they don't show up as spurious variants.",
    )
    prep_keep_tags = st.checkbox(
        "↳ …but keep the base text's tags in the golden output as milestones",
        value=True,
        disabled=not prep_tags,
        help="The stripped tags from the base/golden text are re-inserted "
        "into the downloaded golden document (in italics, at their original "
        "positions) as page references. They are never part of the "
        "collation itself. Tags from the comparison texts are not kept — "
        "the golden document reproduces only the base text.",
    )
    prep_underscore = st.checkbox(
        "Treat _ as a space (EWTS explicit space)",
        value=True,
        help="In EWTS transliteration an underscore marks an explicit space. "
        "Without this, e.g. pa/_bdag and pa/ bdag read as different words.",
    )
    prep_pipe = st.checkbox(
        "Treat | as a shad (EWTS / OCR)",
        value=True,
        help="Some OCR output writes the shad as a pipe. With this on, | "
        "behaves exactly like / — ignored or reported together with shad.",
    )
    prep_head = st.checkbox(
        "Ignore head marks @ # ! (yig-mgo ༄༅)",
        value=True,
        help="These transliterate the ornamental head marks that open a "
        "section; they are structural, not textual, so they never count as "
        "variants.",
    )

ignore_shad = st.checkbox(
    "Ignore shad (།) differences",
    value=True,
    help="When checked, differences that consist only of shad punctuation "
    "(།, ༎, ༔ …) are not reported as variant notes. Uncheck to have shad "
    "differences show up in the apparatus.",
)

apparatus_mode = st.radio(
    "Apparatus type",
    options=["Negative (only variants)", "Positive (all witnesses)"],
    index=0,
    help="Negative apparatus lists only the witnesses that differ from the "
    "base/golden reading. Positive apparatus lists every comparison witness "
    "at each variant point, including those that agree with the lemma.",
)
positive = apparatus_mode.startswith("Positive")

st.subheader("4 · Run")

ready = base_file is not None and comp1_file is not None
if two_texts:
    ready = ready and comp2_file is not None

if not ready:
    st.info("Upload all required files above to enable the collation.")

run_btn = st.button("▶ Run Collation", disabled=not ready, type="primary")

if run_btn and ready:
    # .getvalue() (not .read()) — Streamlit keeps the uploaded file's read
    # cursor across reruns, so a second Run would read empty bytes and
    # silently reuse the previous results.
    text1 = base_file.getvalue().decode("utf-8")
    text2 = comp1_file.getvalue().decode("utf-8")
    text3 = comp2_file.getvalue().decode("utf-8") if comp2_file else ""

    # Apply the user's preprocessing choices
    apply_preprocessing_options(
        underscore_as_space=prep_underscore,
        pipe_as_shad=prep_pipe,
        ignore_head_marks=prep_head,
    )
    prep_preview = {}
    names_texts = [(base_file.name, text1), (comp1_file.name, text2)]
    if comp2_file:
        names_texts.append((comp2_file.name, text3))
    for fname, ftext in names_texts:
        prep_preview[fname] = count_preprocessing_hits(ftext)
    golden_milestones = None
    if prep_tags:
        if prep_keep_tags:
            text1, golden_milestones = extract_folio_tags(text1)
        else:
            text1 = strip_folio_tags(text1)
        text2 = strip_folio_tags(text2)
        text3 = strip_folio_tags(text3) if text3 else text3

    with st.expander("Preprocessing preview", expanded=False):
        st.caption(
            "Occurrences of each preprocessing pattern found per file "
            "(counted before cleanup was applied)."
        )
        for fname, counts in prep_preview.items():
            hits = ", ".join(f"{k}: {v}" for k, v in counts.items() if v) or "nothing to clean"
            st.markdown(f"- **{fname}** — {hits}")

    with st.spinner("Aligning texts… this may take a minute for long texts."):
        if two_texts:
            aligned1, aligned2, aligned3 = align_three(text1, text2, text3)
        else:
            aligned1, aligned2, aligned3 = align_two(text1, text2)

    with st.spinner("Building collation report…"):
        name1 = base_file.name
        name2 = comp1_file.name
        name3 = comp2_file.name if comp2_file else "—"

        report_buf, notes = export_three_way_with_notes(
            aligned1, aligned2, aligned3,
            label1, label2, label3,
            name1=name1, name2=name2, name3=name3,
            ignore_shad=ignore_shad,
            positive=positive,
        )

    footnote_buf = None
    try:
        with st.spinner("Building footnote document…"):
            footnote_buf = export_golden_with_footnotes(
                aligned1, aligned2, aligned3,
                notes,
                label1, label2, label3,
                name1=name1,
                ignore_shad=ignore_shad,
                milestones=golden_milestones,
            )
    except AttributeError:
        st.warning(
            "⚠️ Footnote document skipped — `add_footnote` not available. "
            "Install **bayoo-docx** or **python-docx-2023** instead of python-docx."
        )

    # Store results in session_state so downloads persist after button clicks
    st.session_state["report_buf"] = report_buf.getvalue()
    st.session_state["footnote_buf"] = footnote_buf.getvalue() if footnote_buf else None
    st.session_state["note_count"] = len(notes)

# Show download section whenever results are available in session_state
if "report_buf" in st.session_state:
    st.success(f"Done! Found **{st.session_state['note_count']}** difference note(s).")
    st.divider()
    st.subheader("5 · Download outputs")

    dl1, dl2 = st.columns(2)
    with dl1:
        st.download_button(
            label="⬇ Collation report (.docx)",
            data=st.session_state["report_buf"],
            file_name="collation_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_report",
        )
    with dl2:
        if st.session_state["footnote_buf"]:
            st.download_button(
                label="⬇ Golden text + footnotes (.docx)",
                data=st.session_state["footnote_buf"],
                file_name="collation_footnotes.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_footnotes",
            )
        else:
            st.button("⬇ Golden text + footnotes (.docx)", disabled=True)

st.divider()
REPO_URL = "https://github.com/jyerena108/tibetan-collation"

with st.expander("ℹ️ About this tool"):
    st.markdown(f"""
**Tibetan Collation Tool** — aligns two or three versions of a Tibetan text
and generates a critical apparatus in Word format. Works with Unicode Tibetan
and Wylie/EWTS transliteration.

Notes read `BX1 kyi] AB1, GB1 ni` — *where BX1 reads `kyi`, AB1 and GB1 read
`ni`*. Omissions are marked `om.`

**How to cite**

> Yerena, J. *Tibetan Collation Tool* (2026). {REPO_URL}

If the alignment matters to your argument, please also cite Pydurma:

> Roux, E. & Kaldan, T. *Pydurma* (2023). https://github.com/openpecha/pydurma

**Run your own copy**

```bash
git clone {REPO_URL}.git
cd tibetan-collation
pip install -r requirements.txt
streamlit run collation_app_01.py
```

Full documentation is in the [README]({REPO_URL}#readme).
""")

st.caption(
    f"**Tibetan Collation Tool** by Yerena, J. — free and open source "
    f"([MIT]({REPO_URL}/blob/main/LICENSE)) · [source code]({REPO_URL})  \n"
    "Alignment by [Pydurma](https://github.com/openpecha/pydurma) "
    "(OpenPecha, MIT).  \n"
    "Provided as-is, without warranty. Please check collation results before "
    "relying on them in published work."
)
